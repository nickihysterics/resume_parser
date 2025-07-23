import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """
    Извлекает текст из PDF-файла постранично.
    Игнорирует страницы без текста.
    """
    with pdfplumber.open(filepath) as pdf:
        return '\n'.join(page.extract_text() or '' for page in pdf.pages)


def extract_text_from_docx(filepath: str) -> str:
    """
    Извлекает текст из DOCX-файла, включая параграфы и таблицы.
    """
    doc = Document(filepath)
    lines = []

    # Основной текст
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    # Текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        lines.append(p.text.strip())

    return '\n'.join(lines)


def extract_text(filepath: str) -> str:
    """
    Определяет тип файла и вызывает соответствующий метод извлечения текста.
    Поддерживаемые форматы: .pdf, .docx
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
